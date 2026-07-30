"""Generate binary fixtures for end-to-end tests from the plain-text sources."""

from __future__ import annotations

import csv
from pathlib import Path

from docx import Document
from docx.shared import Pt
from openpyxl import Workbook
from openpyxl.styles import Font
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

FIXTURES = Path(__file__).parent


def _read_text(name: str) -> str:
    return (FIXTURES / name).read_text(encoding="utf-8")


def _clean_paragraphs(text: str) -> list[str]:
    """Split markdown into paragraphs; keep table rows in one paragraph."""
    out: list[str] = []
    for raw in text.split("\n\n"):
        para = " ".join(line.strip() for line in raw.splitlines() if line.strip())
        if para:
            out.append(para)
    return out


def generate_docx(output: Path, *parts: str) -> None:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)

    for text in parts:
        for para in _clean_paragraphs(text):
            if para.startswith("# "):
                doc.add_heading(para[2:], level=1)
            elif para.startswith("## "):
                doc.add_heading(para[3:], level=2)
            elif para.startswith("##Q"):
                doc.add_paragraph(para)
            else:
                doc.add_paragraph(para)
    doc.save(output)


def generate_pdf(output: Path, *parts: str) -> None:
    doc = SimpleDocTemplate(str(output), pagesize=A4)
    styles = getSampleStyleSheet()
    story: list = []
    for text in parts:
        for para in _clean_paragraphs(text):
            if para.startswith("# "):
                story.append(Paragraph(para[2:], styles["Heading1"]))
            elif para.startswith("## "):
                story.append(Paragraph(para[3:], styles["Heading2"]))
            else:
                story.append(Paragraph(para, styles["BodyText"]))
            story.append(Spacer(1, 6))
    doc.build(story)


def generate_xlsx(csv_path: Path, output: Path) -> None:
    wb = Workbook()
    ws = wb.active
    if ws is None:
        return
    ws.title = "BOQ"

    with csv_path.open(encoding="utf-8") as f:
        reader = csv.reader(f)
        for row in reader:
            ws.append(row)

    # Header style
    for cell in ws[1]:
        cell.font = Font(bold=True)

    # Add a sheet with one intentionally wrong total for defect detection tests
    ws2 = wb.create_sheet("WRONG_TOTAL")
    ws2.append(["Description", "Computed Total", "Stated Total", "Expected Error"])
    ws2.append(["RCC M25 in columns", "1062500", "1162500", "Overstated by 100000"])
    ws2.append(["Vitrified tile flooring", "1140000", "1040000", "Understated by 100000"])

    wb.save(output)


def main() -> None:
    nit = _read_text("sample_nit.md")
    gcc = _read_text("sample_gcc.md")
    qa = _read_text("sample_pre_bid_qa.md")
    addendum = _read_text("sample_addendum.md")

    tender_pack = "\n\n".join([nit, gcc, addendum])
    generate_docx(FIXTURES / "sample_tender_pack.docx", tender_pack)
    generate_pdf(FIXTURES / "sample_tender_pack.pdf", tender_pack)

    generate_xlsx(FIXTURES / "sample_boq.csv", FIXTURES / "sample_boq.xlsx")

    # A standalone PDF for risk-pattern tests
    generate_pdf(FIXTURES / "sample_gcc.pdf", gcc)

    # A standalone PDF for Q&A tests
    generate_pdf(FIXTURES / "sample_pre_bid_qa.pdf", qa)

    print("Generated fixtures:")
    for f in FIXTURES.iterdir():
        if f.suffix in {".docx", ".pdf", ".xlsx"}:
            print(f"  {f.name}: {f.stat().st_size} bytes")


if __name__ == "__main__":
    main()
