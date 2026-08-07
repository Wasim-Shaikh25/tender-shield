"""Pure renderers for the Bid Review Pack (Doc §1.1(8), §11.4). No DB — take
plain data, return file bytes. Every export carries the review/date/pack stamp."""

from __future__ import annotations

import io

from docx import Document
from openpyxl import Workbook, load_workbook

_SEV_RANK = {"critical": 0, "high": 1, "medium": 2, "low": 3, "info": 4}

UNREVIEWED_VARIANT = "unreviewed"
UNREVIEWED_WATERMARK = "INDICATIVE — NOT REVIEWED BY A QUALIFIED PROFESSIONAL"


def _template_text(template: dict | None, key: str, default: str) -> str:
    return (template or {}).get(key) or default


def stamp_line(meta: dict, template: dict | None = None) -> str:
    footer = _template_text(template, "footer_text", "")
    if meta.get("variant") == UNREVIEWED_VARIANT:
        integrity = meta.get("integrity_hash")
        integrity_text = f" · integrity {integrity[:16]}" if integrity else ""
        base = (
            f"{UNREVIEWED_WATERMARK} · Prepared with TenderShield · "
            f"{meta.get('date', '')} · pack {meta.get('pack', 'in-works')}"
            f"{integrity_text} · Machine-generated output — requires professional review."
        )
    else:
        reviewed = meta.get("reviewed_by_email")
        reviewed_at = meta.get("reviewed_at", "")
        if reviewed:
            reviewer = f" · reviewed by {reviewed} on {reviewed_at}"
        else:
            reviewer = ""
        integrity = meta.get("integrity_hash")
        integrity_text = f" · integrity {integrity[:16]}" if integrity else ""
        base = (
            f"Prepared with TenderShield · reviewed and approved on {meta.get('date', '')} "
            f"· pack {meta.get('pack', 'in-works')}{reviewer}{integrity_text} · This is "
            f"document-intelligence software, not legal/QS advice — review with a qualified "
            f"professional."
        )
    return f"{base} {footer}".strip() if footer else base


def _pdf_page_watermark(canvas, _doc) -> None:
    from reportlab.lib.pagesizes import A4

    canvas.saveState()
    canvas.setFont("Helvetica-Bold", 28)
    canvas.setFillColorRGB(0.75, 0.75, 0.75, alpha=0.35)
    canvas.translate(A4[0] / 2, A4[1] / 2)
    canvas.rotate(45)
    canvas.drawCentredString(0, 0, UNREVIEWED_WATERMARK)
    canvas.restoreState()


def _docx_apply_watermark(doc: Document) -> None:
    for section in doc.sections:
        header = section.header
        para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        para.text = UNREVIEWED_WATERMARK
        if para.runs:
            para.runs[0].bold = True


def verify_unreviewed_watermark(fmt: str, data: bytes) -> bool:
    """Return True when the rendered export carries the unreviewed watermark."""
    if fmt == "xlsx":
        wb = load_workbook(io.BytesIO(data), read_only=True)
        ws = wb.active
        for row in ws.iter_rows(min_row=1, max_row=5, min_col=1, max_col=1):
            for cell in row:
                if cell.value and UNREVIEWED_WATERMARK in str(cell.value):
                    return True
        return False
    if fmt == "docx":
        doc = Document(io.BytesIO(data))
        header = doc.sections[0].header.paragraphs[0].text if doc.sections else ""
        body = "\n".join(p.text for p in doc.paragraphs[:3])
        return UNREVIEWED_WATERMARK in header or UNREVIEWED_WATERMARK in body
    return True


def render_xlsx(
    opportunity_title: str, findings: list[dict], meta: dict, template: dict | None = None
) -> bytes:
    wb = Workbook()
    ws = wb.active
    ws.title = "Risk Register"
    title = _template_text(template, "report_title", "Bid Review Pack")
    ws.append([f"{title} — {opportunity_title}"])
    if meta.get("variant") == UNREVIEWED_VARIANT:
        ws.append([UNREVIEWED_WATERMARK])
    ws.append([stamp_line(meta, template)])
    ws.append([])
    ws.append(["Severity", "Category", "Title", "Review", "Page", "Source quote"])
    for f in sorted(findings, key=lambda x: _SEV_RANK.get(x.get("severity", "info"), 9)):
        ws.append(
            [
                f.get("severity"),
                f.get("category"),
                f.get("title"),
                f.get("review_status"),
                f.get("source_page"),
                f.get("source_quote"),
            ]
        )
    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


def render_docx(
    opportunity_title: str,
    artifacts: list[dict],
    findings: list[dict],
    meta: dict,
    template: dict | None = None,
) -> bytes:
    doc = Document()
    if meta.get("variant") == UNREVIEWED_VARIANT:
        _docx_apply_watermark(doc)
    title = _template_text(template, "report_title", "Bid Review Pack")
    doc.add_heading(f"{title} — {opportunity_title}", level=0)
    doc.add_paragraph(stamp_line(meta, template)).italic = True

    unreviewed = meta.get("variant") == UNREVIEWED_VARIANT
    accepted = findings if unreviewed else [
        f for f in findings if f.get("review_status") in ("accepted", "edited")
    ]
    doc.add_heading("Risk findings", level=1)
    if accepted:
        table = doc.add_table(rows=1, cols=3)
        hdr = table.rows[0].cells
        hdr[0].text, hdr[1].text, hdr[2].text = "Severity", "Category", "Finding"
        for f in accepted:
            row = table.add_row().cells
            row[0].text = str(f.get("severity", ""))
            row[1].text = str(f.get("category", ""))
            row[2].text = str(f.get("title", ""))
    else:
        doc.add_paragraph("No accepted findings.")

    for art in artifacts:
        body = art.get("body", {})
        doc.add_heading(body.get("title", art.get("kind", "Artifact")), level=1)
        if body.get("preamble"):
            doc.add_paragraph(body["preamble"])
        for item in body.get("items", []):
            if art.get("kind") == "clarification_letter":
                doc.add_paragraph(f"{item.get('n')}. {item.get('heading', '')}", style=None)
                if item.get("quote"):
                    doc.add_paragraph(f"“{item['quote']}”").italic = True
                doc.add_paragraph(str(item.get("ask", "")))
            else:
                cat, txt = item.get("category"), item.get("assumption", "")
                doc.add_paragraph(f"{item.get('n')}. [{cat}] {txt}")

    if template and template.get("footer_text"):
        doc.add_paragraph(template["footer_text"])

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def render_pdf(
    opportunity_title: str,
    artifacts: list[dict],
    findings: list[dict],
    meta: dict,
    template: dict | None = None,
) -> bytes:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, PageBreak

    styles = getSampleStyleSheet()
    normal = styles["Normal"]
    small = ParagraphStyle("stamp", parent=normal, fontSize=8, textColor="#666666")
    quote_style = ParagraphStyle("quote", parent=normal, fontSize=9, textColor="#555555", leftIndent=20, italic=True)

    buf = io.BytesIO()
    unreviewed = meta.get("variant") == UNREVIEWED_VARIANT
    page_callbacks = (
        {"onFirstPage": _pdf_page_watermark, "onLaterPages": _pdf_page_watermark}
        if unreviewed
        else {}
    )
    title = _template_text(template, "report_title", "Bid Review Pack")
    doc = SimpleDocTemplate(
        buf, pagesize=A4, title=f"{title} — {opportunity_title}", **page_callbacks
    )

    # Cover page
    flow = [
        Paragraph(f"{title}", styles["Title"]),
        Paragraph(f"{opportunity_title}", styles["Heading1"]),
        Spacer(1, 24),
        Paragraph(stamp_line(meta, template), small),
        PageBreak(),
    ]

    # Executive summary
    accepted = findings if unreviewed else [
        f for f in findings if f.get("review_status") in ("accepted", "edited")
    ]

    high_risk = [f for f in accepted if f.get("severity") in ("critical", "high")]
    medium_risk = [f for f in accepted if f.get("severity") == "medium"]
    boq_issues = [f for f in accepted if f.get("category") == "boq"]

    flow.append(Paragraph("Executive Summary", styles["Heading1"]))
    flow.append(Spacer(1, 10))
    flow.append(Paragraph(f"Total findings: <b>{len(accepted)}</b>", normal))
    if high_risk:
        flow.append(Paragraph(f"<b>Critical/High Risk:</b> {len(high_risk)}", normal))
    if medium_risk:
        flow.append(Paragraph(f"<b>Medium Risk:</b> {len(medium_risk)}", normal))
    if boq_issues:
        flow.append(Paragraph(f"<b>BOQ Issues:</b> {len(boq_issues)}", normal))
    flow.append(PageBreak())

    # Risk findings with source quotes
    flow.append(Paragraph("Risk Findings", styles["Heading1"]))
    flow.append(Spacer(1, 10))

    if accepted:
        for severity_level in ("critical", "high", "medium", "low", "info"):
            severity_findings = [f for f in accepted if f.get("severity") == severity_level]
            if severity_findings:
                flow.append(Paragraph(f"<b>{severity_level.upper()} SEVERITY</b>", styles["Heading2"]))
                for f in severity_findings:
                    flow.append(Paragraph(
                        f"<b>{f.get('category', 'unknown').upper()}</b> -- {f.get('title', 'Untitled')}",
                        normal
                    ))
                    quote = f.get("source_quote")
                    if quote:
                        page_ref = f.get("source_page")
                        page_text = f" (Page {page_ref})" if page_ref else ""
                        flow.append(Paragraph(f'<i>"{quote}"</i>{page_text}', quote_style))
                    flow.append(Spacer(1, 6))
                flow.append(Spacer(1, 12))
    else:
        flow.append(Paragraph("No accepted findings.", normal))

    flow.append(PageBreak())

    # Artifacts: Assumptions & Clarifications
    for art in artifacts:
        b = art.get("body", {})
        flow.append(Paragraph(b.get("title", art.get("kind", "Artifact")).upper(), styles["Heading1"]))
        if b.get("preamble"):
            flow.append(Paragraph(b["preamble"], normal))
            flow.append(Spacer(1, 10))

        for item in b.get("items", []):
            if art.get("kind") == "clarification_letter":
                flow.append(Paragraph(f"<b>{item.get('n')}. {item.get('heading', '')}</b>", normal))
                if item.get("quote"):
                    flow.append(Paragraph(f'<i>"{item["quote"]}"</i>', quote_style))
                flow.append(Paragraph(str(item.get("ask", "")), normal))
            else:
                cat, txt = item.get("category"), item.get("assumption", "")
                flow.append(Paragraph(f"<b>{item.get('n')}. [{cat}]</b> {txt}", normal))
            flow.append(Spacer(1, 8))

        flow.append(Spacer(1, 16))

    doc.build(flow)
    return buf.getvalue()

def _handover_table_sections(handover: dict) -> list[tuple[str, list[str], list, list[str]]]:
    specs = [
        (
            "watchlist_controls",
            "Watchlist Controls",
            ["Title", "Severity", "Cadence"],
            ["title", "severity", "review_cadence"],
        ),
        (
            "key_obligations",
            "Key Obligations",
            ["Severity", "Category", "Title", "Quote"],
            ["severity", "category", "title", "source_quote"],
        ),
        (
            "notice_register",
            "Notice Register",
            ["Category", "Days", "Trigger", "Quote"],
            ["category", "days", "trigger", "source_quote"],
        ),
        (
            "notice_gaps",
            "Notice Gaps",
            ["Category", "Typical Days", "Note"],
            ["key", "typical_days", "note"],
        ),
        (
            "deadline_calendar",
            "Deadlines",
            ["Kind", "Due At", "Description"],
            ["kind", "due_at", "description"],
        ),
        (
            "boq_assumptions",
            "BOQ Assumptions",
            ["Category", "Title", "Exposure"],
            ["category", "title", "amount_exposure"],
        ),
        (
            "scope_gap_findings",
            "Scope Gaps",
            ["Category", "Title", "Quote"],
            ["category", "title", "source_quote"],
        ),
        (
            "finance_notice_windows",
            "Finance Notice Windows",
            ["Category", "Days", "Trigger"],
            ["category", "days", "trigger"],
        ),
    ]
    out = []
    for key, heading, headers, fields in specs:
        if key not in handover:
            continue
        out.append((heading, headers, handover.get(key) or [], fields))
    return out


def render_handover_pack(
    opportunity_title: str,
    handover: dict,
    fmt: str,
    meta: dict,
    template: dict | None = None,
) -> bytes:
    """Render a sealed-baseline handover pack in the requested office format."""
    view = handover.get("view") or meta.get("view") or "full"
    sections = _handover_table_sections(handover)
    cost_codes = handover.get("cost_codes") or []
    exposure = handover.get("exposure_totals") or {}

    title = _template_text(template, "report_title", "Handover Pack")
    if fmt == "xlsx":
        wb = Workbook()
        ws = wb.active
        ws.title = "Handover"
        ws.append([f"{title} — {opportunity_title} ({view})"])
        ws.append([stamp_line(meta, template)])
        version = handover.get("version")
        source = handover.get("source")
        sealed_hash = handover.get("sealed_hash", "")[:16]
        ws.append([f"Version: {version}  Source: {source}  Hash: {sealed_hash}"])
        ws.append([])
        for heading, headers, rows, fields in sections:
            ws.append([heading])
            ws.append(headers)
            for r in rows:
                ws.append([r.get(f) for f in fields])
            ws.append([])
        if cost_codes:
            ws.append(["Cost Codes"])
            ws.append(["Code", "Label", "Variation Category", "Mappings"])
            for r in cost_codes:
                ws.append(
                    [
                        r.get("code"),
                        r.get("label"),
                        r.get("variation_category"),
                        len(r.get("mappings", [])),
                    ]
                )
            ws.append([])
        if exposure:
            ws.append(["Exposure Totals"])
            ws.append(
                ["Currency", exposure.get("currency"), "Total Minor", exposure.get("total_minor")]
            )
        buf = io.BytesIO()
        wb.save(buf)
        return buf.getvalue()

    if fmt == "docx":
        doc = Document()
        doc.add_heading(f"{title} — {opportunity_title} ({view})", level=0)
        doc.add_paragraph(stamp_line(meta, template)).italic = True
        doc.add_paragraph(
            f"Version: {handover.get('version')}  Source: {handover.get('source')}  "
            f"Hash: {handover.get('sealed_hash', '')[:16]}"
        )

        for heading, headers, rows, fields in sections:
            doc.add_heading(heading, level=1)
            if not rows:
                doc.add_paragraph("None.")
                continue
            table = doc.add_table(rows=1, cols=len(headers))
            for i, h in enumerate(headers):
                table.rows[0].cells[i].text = h
            for r in rows:
                cells = table.add_row().cells
                for i, f in enumerate(fields):
                    cells[i].text = str(r.get(f))

        if cost_codes:
            doc.add_heading("Cost Codes", level=1)
            table = doc.add_table(rows=1, cols=4)
            for i, h in enumerate(["Code", "Label", "Variation Category", "Mappings"]):
                table.rows[0].cells[i].text = h
            for r in cost_codes:
                cells = table.add_row().cells
                cells[0].text = str(r.get("code"))
                cells[1].text = str(r.get("label"))
                cells[2].text = str(r.get("variation_category"))
                cells[3].text = str(len(r.get("mappings", [])))

        if exposure:
            doc.add_heading("Exposure Totals", level=1)
            doc.add_paragraph(
                f"Currency: {exposure.get('currency')}  "
                f"Total (minor units): {exposure.get('total_minor')}"
            )

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    # pdf
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

    styles = getSampleStyleSheet()
    normal = styles["Normal"]
    small = ParagraphStyle("stamp", parent=normal, fontSize=8, textColor="#666666")
    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4, title=f"{title} — {opportunity_title}")
    flow = [
        Paragraph(f"{title} — {opportunity_title} ({view})", styles["Title"]),
        Paragraph(stamp_line(meta, template), small),
        Paragraph(
            f"Version: {handover.get('version')}  Source: {handover.get('source')}  "
            f"Hash: {handover.get('sealed_hash', '')[:16]}",
            normal,
        ),
        Spacer(1, 12),
    ]

    def pdf_section(heading, rows, formatter):
        flow.append(Spacer(1, 10))
        flow.append(Paragraph(f"<b>{heading}</b>", normal))
        if not rows:
            flow.append(Paragraph("None.", normal))
            return
        for row in rows:
            flow.append(Paragraph(formatter(row), normal))

    for heading, _headers, rows, fields in sections:
        pdf_section(
            heading,
            rows,
            lambda r, fields=fields: " | ".join(
                f"{f}: {r.get(f)}" for f in fields if r.get(f) is not None
            ),
        )
    if exposure:
        summary = (
            f"Currency: {exposure.get('currency')} Total: {exposure.get('total_minor')}"
        )
        pdf_section(
            "Exposure Totals",
            [{"summary": summary}],
            lambda r: r["summary"],
        )
    doc.build(flow)
    return buf.getvalue()
