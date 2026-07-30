"""TenderShield Office MCP server.

Exposes tools for QS engineers to read and update Microsoft Word/Excel tender
documents from an LLM client (Claude Desktop, Cursor, etc.) that speaks the
Model Context Protocol (MCP).

Local file tools read and write Word/Excel files in the current working
directory. TenderShield API tools pull live opportunities, findings, deadlines,
and AI-generated plan dashboards when `TENDERSHIELD_API_BASE` and
`TENDERSHIELD_API_TOKEN` are set.

Transport: stdio (default). Start with:

    python -m tendershield_office_mcp.server

or install as a package and run:

    python -m tendershield_office_mcp
"""

from __future__ import annotations

import json
import os
from pathlib import Path
from typing import Any

from mcp.server.fastmcp import FastMCP
from openpyxl import Workbook, load_workbook

from tendershield_office_mcp.client import TenderShieldClient

try:
    from docx import Document
except ImportError:  # pragma: no cover - guard for typing in minimal envs
    Document = None  # type: ignore[assignment,misc]

mcp = FastMCP("tendershield-office")


def _api_client() -> TenderShieldClient:
    return TenderShieldClient()


def _resolve_path(path: str) -> Path:
    """Reject parent-directory traversal and expand user paths."""
    p = Path(path).expanduser().resolve()
    cwd = Path(os.getcwd()).resolve()
    if not str(p).startswith(str(cwd)):
        raise ValueError("path must be inside the current working directory")
    return p


def _load_docx(path: Path) -> Any:
    if Document is None:
        raise RuntimeError("python-docx is not installed")
    return Document(str(path))


@mcp.tool()
def read_word_document(path: str) -> str:
    """Read the text and tables of a Microsoft Word (.docx) tender document."""
    p = _resolve_path(path)
    if not p.exists():
        return f"File not found: {p}"
    doc = _load_docx(p)
    paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
    tables: list[list[list[str]]] = []
    for table in doc.tables:
        rows = []
        for row in table.rows:
            rows.append([cell.text for cell in row.cells])
        tables.append(rows)
    return json.dumps({"paragraphs": paragraphs, "tables": tables}, indent=2)


@mcp.tool()
def read_excel_workbook(path: str, sheet: str | None = None, max_rows: int = 100) -> str:
    """Read an Excel workbook (or a single sheet) as JSON."""
    p = _resolve_path(path)
    if not p.exists():
        return f"File not found: {p}"
    wb = load_workbook(str(p), data_only=True)
    if sheet:
        sheets = [wb[sheet]]
    else:
        sheets = wb.worksheets
    output: dict[str, list[list[Any]]] = {}
    for ws in sheets:
        rows = []
        for row in ws.iter_rows(values_only=True, max_row=max_rows):
            rows.append([str(v) if v is not None else "" for v in row])
        output[ws.title] = rows
    return json.dumps(output, indent=2)


@mcp.tool()
def write_word_comments(path: str, comments: list[str], heading: str = "TenderShield Review Comments") -> str:
    """Append a section of review comments to an existing Word document."""
    p = _resolve_path(path)
    if not p.exists():
        return f"File not found: {p}"
    doc = _load_docx(p)
    doc.add_heading(heading, level=1)
    for comment in comments:
        doc.add_paragraph(comment, style="List Bullet")
    doc.save(str(p))
    return f"Saved {len(comments)} comments to {p}"


@mcp.tool()
def create_summary_doc(path: str, title: str, sections: list[dict[str, Any]]) -> str:
    """Create a new Word document summarizing tender review findings.

    sections: list of {"heading": str, "items": [str]}
    """
    p = _resolve_path(path)
    if p.exists():
        return f"File already exists: {p}"
    doc = _load_docx_with_new()
    doc.add_heading(title, level=0)
    for section in sections:
        heading = section.get("heading", "Section")
        items = section.get("items", []) or []
        doc.add_heading(heading, level=1)
        for item in items:
            doc.add_paragraph(str(item), style="List Bullet")
    doc.save(str(p))
    return f"Created {p}"


@mcp.tool()
def append_excel_rows(path: str, sheet: str, rows: list[list[Any]]) -> str:
    """Append rows to an Excel worksheet, creating the workbook/sheet if missing."""
    p = _resolve_path(path)
    if p.exists():
        wb = load_workbook(str(p))
        if sheet in wb.sheetnames:
            ws = wb[sheet]
        else:
            ws = wb.create_sheet(title=sheet)
    else:
        wb = Workbook()
        ws = wb.active
        if ws is None:
            ws = wb.create_sheet()
        ws.title = sheet
    for row in rows:
        ws.append([str(v) if v is not None else "" for v in row])
    wb.save(str(p))
    return f"Appended {len(rows)} rows to {p}::{sheet}"


def _load_docx_with_new() -> Any:
    if Document is None:
        raise RuntimeError("python-docx is not installed")
    return Document()


# ---------------------------------------------------------------------------
# TenderShield API-aware tools (TS-190)
# ---------------------------------------------------------------------------


@mcp.tool()
def tendershield_list_opportunities() -> str:
    """List tender opportunities the API token has access to."""
    client = _api_client()
    if not client.configured():
        return (
            "TenderShield API not configured. Set TENDERSHIELD_API_BASE and "
            "TENDERSHIELD_API_TOKEN environment variables."
        )
    try:
        with client:
            return json.dumps(client.list_opportunities(), indent=2, default=str)
    except Exception as exc:  # pragma: no cover - real network errors in MCP
        return f"TenderShield API error: {exc}"


@mcp.tool()
def tendershield_get_opportunity_summary(opportunity_id: str) -> str:
    """Return high-level metadata for a tender opportunity."""
    client = _api_client()
    if not client.configured():
        return (
            "TenderShield API not configured. Set TENDERSHIELD_API_BASE and "
            "TENDERSHIELD_API_TOKEN environment variables."
        )
    try:
        with client:
            opp = client.get_opportunity(opportunity_id)
            deadlines = client.get_deadlines(opportunity_id)
            return json.dumps(
                {
                    "opportunity": opp,
                    "deadline_count": len(deadlines),
                    "deadlines": deadlines[:10],
                },
                indent=2,
                default=str,
            )
    except Exception as exc:  # pragma: no cover
        return f"TenderShield API error: {exc}"


@mcp.tool()
def tendershield_get_findings(opportunity_id: str, severity: str = "") -> str:
    """List risk/BOQ findings for a tender opportunity, optionally filtered by severity."""
    client = _api_client()
    if not client.configured():
        return (
            "TenderShield API not configured. Set TENDERSHIELD_API_BASE and "
            "TENDERSHIELD_API_TOKEN environment variables."
        )
    try:
        with client:
            findings = client.get_findings(opportunity_id)
            if severity:
                findings = [f for f in findings if (f.get("severity") or "").lower() == severity.lower()]
            return json.dumps(findings, indent=2, default=str)
    except Exception as exc:  # pragma: no cover
        return f"TenderShield API error: {exc}"


@mcp.tool()
def tendershield_export_findings_to_excel(opportunity_id: str, output_path: str) -> str:
    """Pull findings for an opportunity and append them to an Excel workbook."""
    client = _api_client()
    if not client.configured():
        return (
            "TenderShield API not configured. Set TENDERSHIELD_API_BASE and "
            "TENDERSHIELD_API_TOKEN environment variables."
        )
    try:
        with client:
            findings = client.get_findings(opportunity_id)
        rows = [
            ["ID", "Severity", "Category", "Title", "Producer", "Page", "Status"]
        ]
        for f in findings:
            rows.append(
                [
                    f.get("id", ""),
                    f.get("severity", ""),
                    f.get("category", ""),
                    f.get("title", ""),
                    f.get("producer", ""),
                    f.get("source_page", ""),
                    f.get("status", ""),
                ]
            )
        return append_excel_rows(output_path, f"findings_{opportunity_id}", rows)
    except Exception as exc:  # pragma: no cover
        return f"TenderShield API error: {exc}"


@mcp.tool()
def tendershield_create_summary_doc(opportunity_id: str, output_path: str) -> str:
    """Create a Word document summarising a tender opportunity's findings."""
    client = _api_client()
    if not client.configured():
        return (
            "TenderShield API not configured. Set TENDERSHIELD_API_BASE and "
            "TENDERSHIELD_API_TOKEN environment variables."
        )
    try:
        with client:
            opp = client.get_opportunity(opportunity_id)
            deadlines = client.get_deadlines(opportunity_id)
            findings = client.get_findings(opportunity_id)

        opp_title = (opp or {}).get("title", opportunity_id)
        sections: list[dict[str, Any]] = [
            {"heading": "Key Deadlines", "items": [d.get("title", str(d)) for d in deadlines[:20]]},
        ]
        by_category: dict[str, list[str]] = {}
        for f in findings:
            category = f.get("category", "General")
            by_category.setdefault(category, []).append(f.get("title", str(f)))
        for category, items in by_category.items():
            sections.append({"heading": f"Findings: {category}", "items": items[:30]})

        return create_summary_doc(
            output_path,
            title=f"TenderShield Summary — {opp_title}",
            sections=sections,
        )
    except Exception as exc:  # pragma: no cover
        return f"TenderShield API error: {exc}"


@mcp.tool()
def tendershield_plan_dashboard(opportunity_id: str, query: str) -> str:
    """Generate a dynamic plan dashboard for an opportunity using TenderShield AI."""
    client = _api_client()
    if not client.configured():
        return (
            "TenderShield API not configured. Set TENDERSHIELD_API_BASE and "
            "TENDERSHIELD_API_TOKEN environment variables."
        )
    try:
        with client:
            dashboard = client.get_plan_dashboard(opportunity_id, query)
            return json.dumps(dashboard, indent=2, default=str)
    except Exception as exc:  # pragma: no cover
        return f"TenderShield API error: {exc}"


if __name__ == "__main__":
    mcp.run()
